import tkinter as tk
from tkinter import scrolledtext, ttk, filedialog
from docx import Document
from datetime import datetime
from openai import OpenAI
import threading
import os
import base64
from PIL import Image, ImageTk
import requests
from io import BytesIO
import sv_ttk  # 모던한 UI를 위한 Sun Valley ttk 테마
import time


class MultiAIApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Assistant")
        self.root.geometry("1000x800")

        # 테마 설정
        sv_ttk.set_theme("dark")  # 다크 테마 적용 (light도 가능)

        # 스타일 설정
        self.style = ttk.Style()
        self.style.configure('Chat.TFrame', padding=10)
        self.style.configure('Control.TFrame', padding=5)
        self.style.configure('Title.TLabel', font=('Helvetica', 12, 'bold'))
        self.style.configure('Status.TLabel', font=('Helvetica', 10))

        # OpenAI 클라이언트 및 대화 기록
        self.client = OpenAI()
        self.conversation_history = []

        # UI 초기화 전에 필요한 메서드들 정의
        self.current_photo = None
        self.chat_area = None
        self.message_entry = None
        self.loading_label = None
        self.image_prompt = None
        self.image_label = None
        self.image_status = None
        self.tts_text = None
        self.stt_result = None

        self.setup_ui()

    def send_message(self, event=None):
        message = self.message_entry.get().strip()
        if not message:
            return

        self.message_entry.delete(0, tk.END)
        self.chat_area.insert(tk.END, f"나: {message}\n")
        self.chat_area.see(tk.END)

        self.loading_label.config(text="🤔 응답을 생성하는 중...")
        threading.Thread(target=self.get_gpt_response,
                         args=(message,)).start()

    def get_gpt_response(self, message):
        try:
            self.conversation_history.append({"role": "user", "content": message})

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=self.conversation_history
            )

            assistant_message = response.choices[0].message.content
            self.conversation_history.append(
                {"role": "assistant", "content": assistant_message}
            )

            self.root.after(0, self.update_chat_area, assistant_message)
        except Exception as e:
            self.root.after(0, self.update_chat_area, f"❌ Error: {str(e)}")
        finally:
            self.root.after(0, self.loading_label.config, {"text": ""})

    def update_chat_area(self, response_text):
        self.chat_area.insert(tk.END, f"GPT: {response_text}\n\n")
        self.chat_area.see(tk.END)

    def new_chat(self):
        if self.chat_area.get("1.0", tk.END).strip():
            self.auto_save_chat()

        self.conversation_history = []
        self.chat_area.delete("1.0", tk.END)
        self.chat_area.insert(tk.END, "✨ 새로운 대화가 시작되었습니다.\n\n")
        self.chat_area.see(tk.END)

    def auto_save_chat(self):
        save_dir = "chat_history"
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)

        filename = f"대화내역_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(save_dir, filename)

        self.save_chat_to_file(file_path)
        self.chat_area.insert(tk.END, f"💾 대화 내역이 자동 저장되었습니다: {file_path}\n\n")
        self.chat_area.see(tk.END)

    def save_chat_to_file(self, file_path):
        doc = Document()
        doc.add_heading('AI Assistant 대화 내역', 0)
        doc.add_paragraph(
            f"저장 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}"
        )
        doc.add_paragraph('=' * 50)

        chat_content = self.chat_area.get("1.0", tk.END)
        paragraphs = chat_content.split('\n')

        for para in paragraphs:
            if para.strip():
                if para.startswith('나:'):
                    p = doc.add_paragraph()
                    p.add_run('나: ').bold = True
                    p.add_run(para[3:])
                elif para.startswith('GPT:'):
                    p = doc.add_paragraph()
                    p.add_run('GPT: ').bold = True
                    p.add_run(para[4:])
                else:
                    doc.add_paragraph(para)

        doc.save(file_path)

    def save_to_docx(self):
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                initialfile=f"대화내역_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            if not file_path:
                return

            self.save_chat_to_file(file_path)
            self.chat_area.insert(tk.END, f"💾 대화 내역이 저장되었습니다: {file_path}\n\n")
            self.chat_area.see(tk.END)

        except Exception as e:
            self.chat_area.insert(tk.END,
                                  f"❌ Error: 대화 내역 저장 중 오류가 발생했습니다. {str(e)}\n\n")
            self.chat_area.see(tk.END)

    def setup_ui(self):
        # 메인 컨테이너
        main_container = ttk.Frame(self.root, padding="10")
        main_container.pack(fill='both', expand=True)

        # 노트북(탭) 생성
        self.notebook = ttk.Notebook(main_container)
        self.notebook.pack(fill='both', expand=True)

        # 탭 설정
        self.setup_chat_tab()
        self.setup_image_tab()
        self.setup_voice_tab()

    def setup_chat_tab(self):
        self.chat_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.chat_frame, text=' 💬 채팅 ')

        # 타이틀 프레임
        title_frame = ttk.Frame(self.chat_frame)
        title_frame.pack(fill='x', pady=(0, 10))
        ttk.Label(title_frame, text="AI와의 대화", style='Title.TLabel').pack(side='left')

        # 채팅 영역 (테두리와 패딩 추가)
        chat_container = ttk.Frame(self.chat_frame, relief='solid', borderwidth=1)
        chat_container.pack(fill='both', expand=True, pady=(0, 10))

        self.chat_area = scrolledtext.ScrolledText(
            chat_container,
            wrap=tk.WORD,
            height=20,
            font=('Helvetica', 10),
            bg='#000000'  # 밝은 배경색
        )
        self.chat_area.pack(fill='both', expand=True, padx=5, pady=5)

        # 입력 영역 컨테이너
        input_container = ttk.Frame(self.chat_frame, style='Control.TFrame')
        input_container.pack(fill='x', pady=(0, 5))

        # 메시지 입력 프레임
        input_frame = ttk.Frame(input_container)
        input_frame.pack(fill='x', pady=5)

        self.message_entry = ttk.Entry(input_frame, font=('Helvetica', 10))
        self.message_entry.pack(side='left', fill='x', expand=True, padx=(0, 5))
        self.message_entry.bind('<Return>', self.send_message)

        ttk.Button(
            input_frame,
            text="전송",
            style='Accent.TButton',
            command=self.send_message
        ).pack(side='right')

        # 컨트롤 버튼 프레임
        control_frame = ttk.Frame(input_container)
        control_frame.pack(fill='x', pady=5)

        ttk.Button(
            control_frame,
            text="💾 대화 내역 저장",
            command=self.save_to_docx
        ).pack(side='left', padx=2)

        ttk.Button(
            control_frame,
            text="🔄 새 채팅",
            command=self.new_chat
        ).pack(side='left', padx=2)

        ttk.Button(
            control_frame,
            text="🔊 음성으로 듣기",
            command=self.speak_last_response
        ).pack(side='left', padx=2)

        # 상태 표시 레이블
        self.loading_label = ttk.Label(
            self.chat_frame,
            text="",
            style='Status.TLabel'
        )
        self.loading_label.pack(pady=5)

    def setup_image_tab(self):
        self.image_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.image_frame, text=' 🎨 이미지 생성 ')

        # 타이틀
        ttk.Label(
            self.image_frame,
            text="AI 이미지 생성",
            style='Title.TLabel'
        ).pack(pady=(0, 10))

        # 컨트롤 패널
        control_panel = ttk.LabelFrame(self.image_frame, text="이미지 설정", padding=10)
        control_panel.pack(fill='x', padx=5, pady=5)

        # 프롬프트 입력
        prompt_frame = ttk.Frame(control_panel)
        prompt_frame.pack(fill='x', pady=5)

        ttk.Label(prompt_frame, text="이미지 설명:").pack(side='left', padx=(0, 5))
        self.image_prompt = ttk.Entry(prompt_frame, width=50)
        self.image_prompt.pack(side='left', fill='x', expand=True)

        # 이미지 설정
        settings_frame = ttk.Frame(control_panel)
        settings_frame.pack(fill='x', pady=10)

        # 크기 선택
        size_frame = ttk.LabelFrame(settings_frame, text="크기", padding=5)
        size_frame.pack(side='left', padx=5)

        self.size_var = tk.StringVar(value="1024x1024")
        ttk.Radiobutton(
            size_frame,
            text="정사각형",
            variable=self.size_var,
            value="1024x1024"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            size_frame,
            text="세로형",
            variable=self.size_var,
            value="1024x1792"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            size_frame,
            text="가로형",
            variable=self.size_var,
            value="1792x1024"
        ).pack(side='left', padx=5)

        # 품질 선택
        quality_frame = ttk.LabelFrame(settings_frame, text="품질", padding=5)
        quality_frame.pack(side='left', padx=5)

        self.quality_var = tk.StringVar(value="standard")
        ttk.Radiobutton(
            quality_frame,
            text="일반",
            variable=self.quality_var,
            value="standard"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            quality_frame,
            text="고품질",
            variable=self.quality_var,
            value="hd"
        ).pack(side='left', padx=5)

        # 생성 버튼
        ttk.Button(
            control_panel,
            text="🎨 이미지 생성",
            style='Accent.TButton',
            command=self.generate_image
        ).pack(pady=10)

        # 이미지 표시 영역
        image_display = ttk.LabelFrame(self.image_frame, text="생성된 이미지", padding=10)
        image_display.pack(fill='both', expand=True, padx=5, pady=5)

        self.image_label = ttk.Label(image_display)
        self.image_label.pack(pady=10)

        self.image_status = ttk.Label(
            image_display,
            text="",
            style='Status.TLabel'
        )
        self.image_status.pack(pady=5)

    def setup_voice_tab(self):
        self.voice_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.voice_frame, text=' 🎤 음성 변환 ')

        # TTS 섹션
        tts_frame = ttk.LabelFrame(
            self.voice_frame,
            text="텍스트 → 음성 변환",
            padding=10
        )
        tts_frame.pack(fill='x', padx=5, pady=5)

        ttk.Label(tts_frame, text="변환할 텍스트:").pack(pady=2)

        self.tts_text = scrolledtext.ScrolledText(
            tts_frame,
            height=4,
            font=('Helvetica', 10)
        )
        self.tts_text.pack(fill='x', pady=5)

        # 음성 설정
        voice_frame = ttk.Frame(tts_frame)
        voice_frame.pack(fill='x', pady=5)

        ttk.Label(voice_frame, text="음성 선택:").pack(side='left', padx=5)
        self.voice_var = tk.StringVar(value="alloy")
        voices = ["alloy", "echo", "fable", "onyx", "nova", "shimmer"]
        voice_menu = ttk.OptionMenu(voice_frame, self.voice_var, "alloy", *voices)
        voice_menu.pack(side='left', padx=5)

        ttk.Button(
            tts_frame,
            text="🔊 음성 생성",
            style='Accent.TButton',
            command=self.generate_speech
        ).pack(pady=5)

        # STT 섹션
        stt_frame = ttk.LabelFrame(
            self.voice_frame,
            text="음성 → 텍스트 변환",
            padding=10
        )
        stt_frame.pack(fill='x', padx=5, pady=10)

        # 컨트롤
        control_frame = ttk.Frame(stt_frame)
        control_frame.pack(fill='x', pady=5)

        ttk.Button(
            control_frame,
            text="🎤 음성 파일 선택",
            command=self.select_audio_file
        ).pack(side='left', padx=5)

        # 언어 선택
        self.lang_var = tk.StringVar(value="auto")
        langs = ["auto", "ko", "en", "ja"]
        ttk.Label(control_frame, text="언어:").pack(side='left', padx=5)
        lang_menu = ttk.OptionMenu(control_frame, self.lang_var, "auto", *langs)
        lang_menu.pack(side='left', padx=5)

        # 결과 표시
        self.stt_result = scrolledtext.ScrolledText(
            stt_frame,
            height=4,
            font=('Helvetica', 10)
        )
        self.stt_result.pack(fill='x', pady=5)

    def generate_image(self):
        prompt = self.image_prompt.get().strip()
        if not prompt:
            self.image_status.config(text="⚠️ 이미지 설명을 입력해주세요.")
            return

        self.image_status.config(text="🎨 이미지를 생성하는 중...")
        threading.Thread(target=self.generate_image_thread,
                         args=(prompt,)).start()

    def generate_image_thread(self, prompt):
        try:
            response = self.client.images.generate(
                model="dall-e-3",
                prompt=prompt,
                size=self.size_var.get(),
                quality=self.quality_var.get(),
                n=1,
                style="natural"
            )

            # 이미지 URL 가져오기
            image_url = response.data[0].url
            image_response = requests.get(image_url)
            image = Image.open(BytesIO(image_response.content))

            # 이미지 크기 조정
            max_size = (800, 800)  # 크기 증가
            image.thumbnail(max_size, Image.Resampling.LANCZOS)

            photo = ImageTk.PhotoImage(image)

            # 이미지와 상태 메시지 업데이트
            self.root.after(0, self.update_image, photo, "✨ 이미지가 생성되었습니다!")

            # 이미지 자동 저장
            save_dir = "generated_images"
            if not os.path.exists(save_dir):
                os.makedirs(save_dir)

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            save_path = os.path.join(save_dir, f"generated_image_{timestamp}.png")
            image.save(save_path)

        except Exception as e:
            error_message = str(e)
            self.root.after(0, self.image_status.config,
                            {"text": f"❌ Error: {error_message}"})
            print(f"Error in image generation: {error_message}")

    def update_image(self, photo, status_text):
        self.current_photo = photo  # 참조 유지
        self.image_label.config(image=photo)
        self.image_status.config(text=status_text)

    def generate_speech(self):
        text = self.tts_text.get("1.0", tk.END).strip()
        if not text:
            return

        try:
            response = self.client.audio.speech.create(
                model="tts-1",
                voice=self.voice_var.get(),
                input=text
            )

            # 임시 파일로 저장
            temp_file = "temp_speech.mp3"
            response.stream_to_file(temp_file)

            # 파일 저장 다이얼로그
            save_path = filedialog.asksaveasfilename(
                defaultextension=".mp3",
                filetypes=[("MP3 files", "*.mp3")],
                initialfile=f"speech_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3"
            )

            if save_path:
                # 임시 파일을 선택한 위치로 이동
                if os.path.exists(temp_file):
                    os.replace(temp_file, save_path)

        except Exception as e:
            if os.path.exists(temp_file):
                os.remove(temp_file)
            raise e

    def speak_last_response(self):
        # 마지막 GPT 응답을 음성으로 변환
        chat_content = self.chat_area.get("1.0", tk.END)
        lines = chat_content.split('\n')
        last_response = None

        # 마지막 GPT 응답 찾기
        for line in reversed(lines):
            if line.startswith("GPT: "):
                last_response = line[5:]  # "GPT: " 제거
                break

        if last_response:
            self.tts_text.delete("1.0", tk.END)
            self.tts_text.insert("1.0", last_response)
            self.generate_speech()

    def select_audio_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Audio files", "*.mp3 *.mp4 *.mpeg *.mpga *.m4a *.wav *.webm")
            ]
        )

        if not file_path:
            return

        try:
            with open(file_path, "rb") as audio_file:
                language = self.lang_var.get()
                if language == "auto":
                    language = None

                transcript = self.client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file,
                    language=language
                )

                self.stt_result.delete("1.0", tk.END)
                self.stt_result.insert("1.0", transcript.text)

        except Exception as e:
            self.stt_result.delete("1.0", tk.END)
            self.stt_result.insert("1.0", f"❌ Error: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = MultiAIApp(root)
    root.mainloop()